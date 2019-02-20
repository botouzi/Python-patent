#作者：泊头子 微信公众号：专利方舟
import os
from win32com import client
import docx
from docx import Document #导入库
import xlwt
from xlrd import open_workbook
from xlutils.copy import copy
def doc_to_docx(path,hang):#doc 转成docx
    if os.path.splitext(path)[1] == ".doc":
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(path)  # 目标路径下的文件
        doc.SaveAs(os.path.splitext(path)[0]+".docx", 16)  # 转化后路径下的文件
        doc.Close()
        path1=os.path.splitext(path)[0]+".docx"
        print(path1)
        word.Quit()
        file=docx.Document(path1)
        strN="三、决定"
        strwuxiao="无效"
        for i in range(len(file.paragraphs)):#在该函数中顺便读取无效结论
            if strN in file.paragraphs[i].text:
                if (len(file.paragraphs[i+1].text))<35:
                    if strwuxiao in file.paragraphs[i+1].text:
                        wuxiaojielun="宣告专利权全部无效"
                    else:
                        wuxiaojielun="维持专利权有效"
                else:
                    wuxiaojielun="宣告专利权部分有效"
        document = Document(path1) #读入文件
        tables = document.tables #获取文件中的表格集
        table = tables[1]#获取文件中的第一个表格
        rb = open_workbook("C:\\Python37\\Scripts\\test.xls")
        wb = copy(rb)
        sheet = wb.get_sheet(0)
        futu="图"
        for j in range(0,len(table.rows)):
            sheet.write(hang, j, table.cell(j,1).text)
        sheet.write(hang,len(table.rows), wuxiaojielun)
        wb.save("C:\\Python37\\Scripts\\test.xls")
def find_file(path, ext, file_list=[]):
    dir = os.listdir(path)
    for i in dir:
        i = os.path.join(path, i)
        if os.path.isdir(i):
            find_file(i, ext, file_list)
        else:
            if ext == os.path.splitext(i)[1]:
                file_list.append(i)
    return file_list 
dir_path = "C:\\Python37\\Scripts\\2"#批量转换文件夹
ext = ".doc"
file_list = find_file(dir_path, ext)
hang=1
for file in file_list:
    doc_to_docx(file,hang)
    hang=hang+1
